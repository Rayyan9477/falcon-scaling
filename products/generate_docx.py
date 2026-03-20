"""
Professional DOCX Document Generator
=====================================
Converts all markdown documentation into polished, professional Word documents
using python-docx. Outputs are suitable for client delivery and investor review.

Generates:
  products/docs/approach_guide.docx
  products/docs/product_overview.docx
  products/docs/build_process.docx
  products/docs/custom_gpt_prompt.docx
  research/docs/methodology.docx
  analysis/docs/conversion_analysis.docx
"""

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
BRAND_NAVY = RGBColor(0x1F, 0x4E, 0x79)
BRAND_BLUE = RGBColor(0x2E, 0x75, 0xB6)
BRAND_DARK = RGBColor(0x2C, 0x2C, 0x2C)
BRAND_GRAY = RGBColor(0x66, 0x66, 0x66)
BRAND_LIGHT_BG = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ROOT = Path(__file__).resolve().parent.parent


# ---------------------------------------------------------------------------
# Document styling helpers
# ---------------------------------------------------------------------------

def _setup_doc(title: str, subtitle: str = "", author: str = "Falcon Scaling") -> Document:
    """Create a branded Document with custom styles and a title page."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BRAND_DARK

    # Paragraph spacing
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, (size, color) in enumerate(
        [(24, BRAND_NAVY), (18, BRAND_NAVY), (14, BRAND_BLUE), (12, BRAND_BLUE)],
        start=1,
    ):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = color
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        h_style.paragraph_format.space_after = Pt(6)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Title page ---
    # Add spacing before title
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = Pt(32)
    run.font.color.rgb = BRAND_NAVY
    run.font.bold = True
    run.font.name = "Calibri"

    # Subtitle
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_GRAY
        run.font.name = "Calibri"

    # Divider line
    doc.add_paragraph("")
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("_" * 60)
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(10)

    # Author and date
    doc.add_paragraph("")
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_meta.add_run(f"Prepared by {author}")
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = "Calibri"

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(date.today().strftime("%B %Y"))
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = "Calibri"

    # Page break after title page
    doc.add_page_break()

    # Set core properties
    doc.core_properties.author = author
    doc.core_properties.title = title

    return doc


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading."""
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
              color: RGBColor | None = None, size: int | None = None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def _add_bullet(doc: Document, text: str, level: int = 0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p


def _add_numbered(doc: Document, text: str):
    """Add a numbered list item."""
    return doc.add_paragraph(text, style="List Number")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a professional styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(header)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        # Set cell background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "1F4E79",
            qn("w:val"): "clear",
        })
        shading.append(shading_elm)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            para = cell.paragraphs[0]
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    # Set column widths proportionally
    total = Inches(6)
    col_width = total / len(headers)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    doc.add_paragraph("")  # spacing after table
    return table


def _add_callout(doc: Document, text: str, label: str = "Key Insight"):
    """Add a highlighted callout box."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = BRAND_DARK
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def _add_code_block(doc: Document, text: str):
    """Add a code/template block with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_DARK


# ---------------------------------------------------------------------------
# Document 1: Approach Guide (The Family Office Approach Playbook)
# ---------------------------------------------------------------------------

def build_approach_guide() -> Path:
    """Generate the Family Office Approach Playbook as a professional DOCX."""
    doc = _setup_doc(
        title="The Family Office\nApproach Playbook",
        subtitle="A Practitioner's Guide to Co-Investment with Family Offices",
    )

    # --- Table of Contents placeholder ---
    _add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "Chapter 1: How Family Offices Make Co-Investment Decisions",
        "Chapter 2: The 5 Types of FO Co-Investors",
        "Chapter 3: Conference Strategy",
        "Chapter 4: Outreach Templates",
        "Chapter 5: Red Flags That Get You Blacklisted",
        "Chapter 6: Building Long-Term FO Relationships",
        "Regulatory Considerations",
    ]
    for i, item in enumerate(toc_items, 1):
        _add_para(doc, f"{i}.  {item}", color=BRAND_BLUE)
    doc.add_page_break()

    # --- Chapter 1 ---
    _add_heading(doc, "Chapter 1: How Family Offices Make Co-Investment Decisions", level=1)

    _add_heading(doc, "The Decision Architecture", level=2)
    _add_para(doc, "Family office co-investment decisions follow a fundamentally different process than institutional LP commitments. Understanding this is the difference between getting a meeting and getting ignored.")

    _add_para(doc, "The 3-Gate Model:", bold=True)
    _add_numbered(doc, "Gate 1 — Trust Filter (Eliminates 80% of inbound): Does the FO know you, or someone who knows you? Cold outreach has a <5% response rate with FOs. Warm introductions pass this gate.")
    _add_numbered(doc, "Gate 2 — Thesis Alignment (Eliminates 50% of remaining): Does your deal/fund match their investment thesis? FOs are thesis-driven, not return-driven. A 3x healthcare deal is irrelevant to an FO focused on real estate.")
    _add_numbered(doc, "Gate 3 — Principal Conviction (Final filter): Does the principal (family member or CIO) personally believe in this? FOs don't have investment committees in the institutional sense. One person's conviction often drives the decision.")

    _add_heading(doc, "SFO vs. MFO Decision Dynamics", level=2)
    _add_para(doc, "Single Family Offices:", bold=True)
    _add_para(doc, "Decisions are faster (days to weeks), but harder to access. One person often has final say. Relationships are personal — if the principal likes you, deals happen. If they don't, no amount of returns data matters.")
    _add_para(doc, "Multi-Family Offices:", bold=True)
    _add_para(doc, "More structured process (similar to institutional), multiple stakeholders, formal due diligence. Slower but more predictable. Access is through the professional staff, not the families directly.")

    _add_heading(doc, "What FOs Care About (That Most Fund Managers Miss)", level=2)
    _add_numbered(doc, "Alignment of interest: Do you have significant personal capital in the deal?")
    _add_numbered(doc, "Information flow: Will they get regular, candid updates? Not quarterly reports — real conversations.")
    _add_numbered(doc, "Co-investment rights: Can they participate in follow-on rounds or adjacent deals?")
    _add_numbered(doc, "Reputation risk: Will this deal embarrass the family? FOs are ultra-conservative about public perception.")
    _add_numbered(doc, "Exit timeline: Most FOs think in decades, not fund cycles. Multi-generational time horizons.")

    doc.add_page_break()

    # --- Chapter 2 ---
    _add_heading(doc, "Chapter 2: The 5 Types of FO Co-Investors", level=1)

    types_data = [
        ("Type 1: The Active Allocator (Co-Invest Score 8-10)",
         "Dedicated investment team, regular deal flow review, systematic process.",
         "Iconiq Capital, Bessemer Trust, Cascade Investment",
         "Formal pitch process through their deal intake. They expect institutional-quality materials. Lead with data, not relationships.",
         "1-page teaser, then full deck upon request. Include detailed financials."),
        ("Type 2: The Opportunistic Principal (Co-Invest Score 6-8)",
         "Family principal who does select direct deals based on personal interest/expertise.",
         "Thiel Capital, Schmidt Family Office, Pritzker Group",
         "Through mutual connections. The principal must personally connect with the opportunity. Lead with the story, not the numbers.",
         "Short narrative email (3 paragraphs max), then a conversation, then materials."),
        ("Type 3: The Fund-of-Funds LP (Co-Invest Score 4-6)",
         "Primarily invests through funds, occasionally does co-investments alongside trusted GPs.",
         "Many European and Asian FOs in this category.",
         "Through their existing GP relationships. If you're not already in their fund portfolio, co-investment is unlikely.",
         'Introduction via mutual GP. "We\'re doing a co-invest alongside [GP they know]."'),
        ("Type 4: The Sector Specialist (Co-Invest Score 5-8)",
         "Deep expertise in one sector, only co-invests in that vertical.",
         "Poonawalla (healthcare), Grok Ventures (deep tech), Koch (energy/industrials)",
         "Through sector-specific conferences and networks. Lead with domain expertise.",
         'Sector thesis first, deal second. "Here\'s how we see the [sector] landscape evolving."'),
        ("Type 5: The Legacy Holder (Co-Invest Score 1-3)",
         "Primarily manages existing family assets, minimal new investment activity.",
         "Cadogan Estate, Portman Family Office, many 3rd+ generation offices",
         "Generally not worth pursuing for co-investment. Their mandate is preservation, not growth.",
         "Don't. Unless you have a warm intro and a property/infrastructure deal."),
    ]

    for type_name, profile, examples, approach, what_to_send in types_data:
        _add_heading(doc, type_name, level=2)
        _add_para(doc, f"Profile: {profile}", bold=False)
        _add_para(doc, f"Examples: {examples}", italic=True, color=BRAND_GRAY)
        _add_para(doc, f"How to approach: {approach}")
        _add_para(doc, f"What to send: {what_to_send}")

    doc.add_page_break()

    # --- Chapter 3 ---
    _add_heading(doc, "Chapter 3: Conference Strategy", level=1)

    _add_heading(doc, "Tier 1 Events (Highest FO Density)", level=2)
    _add_table(doc,
        ["Event", "When", "Where", "FO Attendance"],
        [
            ["SuperReturn International", "June", "Berlin", "500+ FOs"],
            ["SALT Conference", "May", "New York", "300+ FOs"],
            ["Milken Institute Conference", "May", "Los Angeles", "200+ FOs"],
            ["Future Investment Initiative (FII)", "October", "Riyadh", "200+ FOs (ME focus)"],
            ["iConnections", "Jan, March", "Miami", "300+ FOs"],
        ],
    )

    _add_heading(doc, "Tier 2 Events (Regional/Specialist)", level=2)
    _add_table(doc,
        ["Event", "Focus"],
        [
            ["Asian Private Banker FO Forum", "Asia-Pacific FOs"],
            ["European Family Office Summit", "European FOs"],
            ["Caproasia Summit", "Southeast Asian FOs"],
            ["Swiss Family Office Forum", "Swiss/DACH FOs"],
            ["Arab Family Office Summit", "Middle East FOs"],
        ],
    )

    _add_heading(doc, "Conference Playbook", level=2)
    _add_numbered(doc, 'Pre-event (2-4 weeks before): Research which FOs are attending. Send targeted LinkedIn connection requests. "Looking forward to connecting at [Event]."')
    _add_numbered(doc, "At the event: Focus on 5-7 target FOs. Don't pitch in the first conversation — build rapport. Ask about their investment thesis, not your deal.")
    _add_numbered(doc, 'Post-event (within 48 hours): Send personalized follow-up referencing your conversation. "Great discussing [specific topic]."')
    _add_numbered(doc, "Follow-through (2-4 weeks later): Share something valuable (research, deal insight, introduction) before asking for anything.")

    doc.add_page_break()

    # --- Chapter 4 ---
    _add_heading(doc, "Chapter 4: Outreach Templates", level=1)

    _add_heading(doc, "Template 1: Warm Introduction Email", level=2)
    _add_code_block(doc,
        "Subject: Introduction — [Your Name], [Your Fund]\n\n"
        "Hi [First Name],\n\n"
        "[Mutual Contact] suggested I reach out. I run [Fund Name], a [$ size] fund\n"
        "focused on [sector].\n\n"
        "I noticed [FO Name]'s investments in [specific sector/companies] — our thesis\n"
        "around [specific angle] aligns closely with your focus.\n\n"
        "Would you be open to a brief call to explore potential synergies? I'm not\n"
        "looking to pitch — genuinely interested in your perspective on [sector trend].\n\n"
        "Best,\n[Your Name]"
    )

    _add_heading(doc, "Template 2: Conference Follow-Up", level=2)
    _add_code_block(doc,
        "Subject: Great connecting at [Event Name]\n\n"
        "Hi [First Name],\n\n"
        "Really enjoyed our conversation about [specific topic] at [Event] on [day].\n\n"
        "Your point about [specific insight they shared] stuck with me — I think\n"
        "there's a meaningful overlap with what we're building at [Fund Name].\n\n"
        "I'd love to continue the conversation. Would [specific date/time] work\n"
        "for a 20-minute call?\n\n"
        "Best,\n[Your Name]"
    )

    _add_heading(doc, "Template 3: LinkedIn Connection Request", level=2)
    _add_code_block(doc,
        "[First Name] — I lead [Fund Name], a [sector]-focused fund. Your work at\n"
        "[FO Name] in [specific area] is impressive. Would love to connect and\n"
        "exchange perspectives on [sector trend]. No pitch — genuine interest."
    )

    doc.add_page_break()

    # --- Chapter 5 ---
    _add_heading(doc, "Chapter 5: Red Flags That Get You Blacklisted", level=1)

    red_flags = [
        ("Mass emails", "FO decision makers can spot batch outreach instantly. One generic email destroys your reputation permanently."),
        ("Name-dropping without permission", '"I was talking to [other FO principal]..." Unless they explicitly authorized you to mention them, don\'t.'),
        ("Inflated track record", "FOs do deep reference checks. If your 3x return was really 1.5x after fees, they'll find out."),
        ("Pushing for a meeting", '"Just following up for the 4th time..." Two follow-ups maximum. If they haven\'t responded, they\'re not interested.'),
        ("Sharing confidential deal info", "If you share another FO's deal data to impress them, they'll assume you'll share theirs too."),
        ("Misunderstanding family dynamics", "Approaching the wrong family member, or not understanding internal family tensions. Research the family structure first."),
        ("Treating FOs like institutional LPs", "Sending 60-page DDQs, requesting formal meetings with 'the investment committee.' FOs don't work like pension funds."),
    ]

    for i, (flag, explanation) in enumerate(red_flags, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {flag}: ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
        run = p.add_run(explanation)

    doc.add_page_break()

    # --- Chapter 6 ---
    _add_heading(doc, "Chapter 6: Building Long-Term FO Relationships", level=1)

    _add_heading(doc, "The 12-Month Relationship Framework", level=2)

    phases = [
        ("Months 1-3: Establish Presence", [
            "Connect on LinkedIn",
            "Share 2-3 pieces of genuinely useful sector research (not your marketing materials)",
            "Attend one event where they'll be present",
        ]),
        ("Months 4-6: Add Value First", [
            "Make an introduction that benefits THEM (not you)",
            "Share a deal they might be interested in (even if it's not yours)",
            "Comment thoughtfully on their LinkedIn posts or press mentions",
        ]),
        ("Months 7-9: Deepen the Relationship", [
            'Request a casual coffee/call — "No agenda, just want to understand your current thinking on [sector]"',
            "Share proprietary market insights from your deal flow",
            "If appropriate, invite them to a private event you're hosting",
        ]),
        ("Months 10-12: Present an Opportunity", [
            "By now they know you, trust you, and understand your expertise",
            "Present your deal/fund as a natural extension of the relationship",
            'Frame it as: "Based on our conversations about [sector], I think this might align with what you\'re looking for"',
        ]),
    ]

    for phase_title, items in phases:
        _add_heading(doc, phase_title, level=3)
        for item in items:
            _add_bullet(doc, item)

    _add_callout(doc, "Give before you ask. Always. The fund managers who succeed with family offices are the ones who spend 12 months building a relationship before ever mentioning a deal.", label="The Cardinal Rule")

    doc.add_page_break()

    # --- Regulatory Considerations ---
    _add_heading(doc, "Important: Regulatory Considerations", level=1)
    _add_para(doc, "This guide provides relationship and networking strategy. It does not constitute legal or compliance advice. Fund managers and placement agents should be aware of:")
    _add_bullet(doc, "US: Regulation D (Rule 506(b) and 506(c)) governs private placement solicitation. General solicitation rules apply.")
    _add_bullet(doc, "Europe: AIFMD marketing passport rules restrict how non-EU managers can approach EU-based FOs.")
    _add_bullet(doc, "Asia: MAS (Singapore), SFC (Hong Kong), and SEBI (India) each have distinct rules.")
    _add_bullet(doc, "Placement agents: May require FINRA registration (US), FCA authorization (UK), or equivalent local licensing.")

    _add_callout(doc, "Always consult qualified legal counsel before initiating any fundraising outreach.", label="Disclaimer")

    _add_para(doc, "Key industry networks for staying current on FO regulatory frameworks:", bold=True)
    _add_bullet(doc, "Family Office Exchange (FOX) — the most established FO peer network globally")
    _add_bullet(doc, "Tiger 21 — peer-to-peer network for UHNW investors and family offices")
    _add_bullet(doc, "CAIA Association — credentialing body for alternative investment professionals")

    # Save
    out = ROOT / "products" / "docs" / "approach_guide.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Document 2: Product Overview
# ---------------------------------------------------------------------------

def build_product_overview() -> Path:
    doc = _setup_doc(
        title="Co-Investment Intelligence\nAccelerator",
        subtitle="Product Overview — $197 One-Time Purchase",
    )

    _add_heading(doc, "The Problem", level=1)
    _add_para(doc, "Fund managers, GPs, and capital raisers waste months trying to figure out which family offices actually co-invest, who they co-invest with, and how to get into co-investment deal flow.")
    _add_callout(doc, "69% of family offices co-invest or plan to (UBS Global Family Office Report 2024). But the 'who invests with whom' network is opaque.", label="Market Reality")
    _add_para(doc, "Current options are either:")
    _add_bullet(doc, "Too expensive: PitchBook ($20K+/yr), Fintrx ($10K+/yr) — out of reach for emerging managers")
    _add_bullet(doc, "Too generic: LinkedIn research, conference networking — unstructured, time-consuming")
    _add_bullet(doc, "Non-existent: No product specifically maps FO co-investment networks at an accessible price point")

    _add_heading(doc, "The Ideal Customer Profile (ICP)", level=1)
    _add_bullet(doc, "Emerging fund managers ($50M-$500M AUM) raising capital from family offices")
    _add_bullet(doc, "PE/VC associates responsible for LP targeting and capital introduction")
    _add_bullet(doc, "Capital introduction consultants building FO relationship maps")
    _add_bullet(doc, "Placement agents working fundraising mandates")
    _add_bullet(doc, "Independent sponsors seeking co-investment partners for deals")

    _add_heading(doc, "The Product ($197 One-Time)", level=1)
    _add_para(doc, "A three-component bundle delivering immediate, actionable co-investment intelligence:")

    _add_heading(doc, 'Component 1: Excel Dashboard — "FO Co-Investment Scorer & Mapper"', level=2)
    _add_bullet(doc, "200 family offices scored for co-investment propensity (1-10 scale)")
    _add_bullet(doc, "Accessibility scoring for outreach readiness (Green/Yellow/Red)")
    _add_bullet(doc, "Known co-investor relationships and fund relationships included per FO")
    _add_bullet(doc, "Filterable by region, sector, check size, co-invest frequency")
    _add_bullet(doc, "Pre-built pivot views by region and sector")

    _add_heading(doc, 'Component 2: Written Guide — "The Family Office Approach Playbook"', level=2)
    _add_bullet(doc, "How family offices make co-investment decisions")
    _add_bullet(doc, "The 5 types of FO co-investors and how to approach each")
    _add_bullet(doc, "Conference strategy guide with specific events")
    _add_bullet(doc, "Email and LinkedIn outreach templates by FO type")
    _add_bullet(doc, "Red flags that get you blacklisted")
    _add_bullet(doc, "Long-term relationship building framework")

    _add_heading(doc, 'Component 3: Custom GPT — "FO Co-Investment Advisor"', level=2)
    _add_bullet(doc, "Interactive AI advisor trained on co-investment methodology")
    _add_bullet(doc, 'Ask: "I\'m raising a $75M climate tech fund — which FOs should I target?"')
    _add_bullet(doc, "Returns personalized targeting lists with reasoning and approach strategy")

    _add_heading(doc, "Ascension Path", level=1)
    _add_table(doc,
        ["Stage", "Product", "Price"],
        [
            ["Entry", "Co-Investment Accelerator (one-time, static snapshot)", "$197"],
            ["Growth", "PolarityIQ Subscription (live platform, 10,000+ FOs)", "$295+/mo"],
            ["Scale", "Custom FO Targeting Report (done-for-you analysis)", "$2,000+"],
            ["Enterprise", "Managed Introduction Service (concierge networking)", "$5,000+"],
        ],
    )
    _add_para(doc, "Why the ascension works: The static Excel creates desire for live data. The guide teaches methodology but doesn't execute it. The product proves ROI at minimal risk.", italic=True, color=BRAND_GRAY)

    _add_heading(doc, "Why Investors Would Pay $197", level=1)
    _add_numbered(doc, "Immediate ROI: One successful FO introduction can generate $1M+ in LP commitments. $197 is trivial.")
    _add_numbered(doc, "Time savings: Replaces weeks of manual LinkedIn research with a pre-scored, filterable database.")
    _add_numbered(doc, "Competitive edge: Most emerging managers rely on warm intros and luck. This is systematic intelligence.")
    _add_numbered(doc, "No comparable product: Nothing in the market maps FO co-investment networks at this price point.")
    _add_numbered(doc, "Risk-free entry: At $197, the purchase decision is a no-brainer for anyone serious about FO fundraising.")

    out = ROOT / "products" / "docs" / "product_overview.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Document 3: Build Process
# ---------------------------------------------------------------------------

def build_build_process() -> Path:
    doc = _setup_doc(
        title="Build Process\nDocumentation",
        subtitle="Co-Investment Intelligence Accelerator ($197) — Step-by-Step",
    )

    _add_heading(doc, "Step 1: Problem Identification & ICP Definition", level=1)
    _add_para(doc, "Tool: Human analysis + AI research  |  Time: 1 hour", italic=True, color=BRAND_GRAY)
    _add_bullet(doc, "Analyzed the Falcon Scaling ecosystem: existing products (Family Office Catalyst, Private Markets Catalyst, Valuation Toolkit) to identify gaps")
    _add_bullet(doc, "Identified co-investment as the highest-value niche: 69% of FOs co-invest (UBS 2024 data), yet no affordable product maps co-investment networks")
    _add_bullet(doc, "Defined ICP: emerging fund managers ($50M-$500M AUM), PE/VC associates, placement agents")
    _add_bullet(doc, "Set price at $197: high enough to signal value, low enough for impulse purchase. Clear ascension to PolarityIQ subscription")

    _add_heading(doc, "Step 2: Scoring Model Design", level=1)
    _add_para(doc, "Tool: Python (pandas, numpy)  |  Time: 30 minutes", italic=True, color=BRAND_GRAY)
    _add_para(doc, "Designed a quantitative scoring framework for each of the 200 family offices:")
    _add_heading(doc, "Co-Investment Score (1-10)", level=2)
    _add_bullet(doc, "Co-Invest Frequency: High=3, Medium=2, Low=1")
    _add_bullet(doc, "Direct Investment: Yes=2, No=0")
    _add_bullet(doc, "LP Status: Active LP=3, Selective LP=2, Minimal=1")
    _add_bullet(doc, "GP/Direct Status: Active GP=2, Emerging=1.5, Direct Only=1, Selective=0.5")
    _add_heading(doc, "Accessibility Score (1-10)", level=2)
    _add_bullet(doc, "Data Confidence: High=3, Medium=2, Low=1")
    _add_bullet(doc, "Website exists: +1  |  Email pattern: +1  |  LinkedIn: +1")
    _add_bullet(doc, "Contact Method: Direct=2, Referral=1.5, Conference=1, Cold=0.5")
    _add_bullet(doc, "Phone: +0.5  |  Conference Attendance: +0.5")

    _add_heading(doc, "Step 3: Excel Dashboard Build", level=1)
    _add_para(doc, "Tool: Python (openpyxl), Excel  |  Time: 1 hour", italic=True, color=BRAND_GRAY)
    _add_numbered(doc, "Read family office dataset (200 records, 45 fields)")
    _add_numbered(doc, "Calculate Co-Investment Score, Accessibility Score, Outreach Readiness for each FO")
    _add_numbered(doc, "Generate main dashboard sheet with conditional formatting (green/yellow/red)")
    _add_numbered(doc, "Generate regional breakdown pivot sheet")
    _add_numbered(doc, "Generate sector analysis sheet")
    _add_numbered(doc, "Include scoring methodology documentation sheet")

    _add_heading(doc, "Step 4: Approach Guide Writing", level=1)
    _add_para(doc, "Tool: Claude AI + domain expertise  |  Time: 2 hours", italic=True, color=BRAND_GRAY)
    _add_para(doc, 'Wrote "The Family Office Approach Playbook" — a 6-chapter guide covering the 3-Gate Model, 5 FO co-investor types, conference strategy, outreach templates, red flags, and a 12-month relationship building framework.')

    _add_heading(doc, "Step 5: Custom GPT System Prompt", level=1)
    _add_para(doc, "Tool: Manual crafting + AI assistance  |  Time: 30 minutes", italic=True, color=BRAND_GRAY)
    _add_para(doc, "Created system instructions for a ChatGPT Custom GPT with personality, knowledge base reference, and example conversations.")

    _add_heading(doc, "Step 6: Product Packaging & Documentation", level=1)
    _add_para(doc, "Tool: Manual  |  Time: 30 minutes", italic=True, color=BRAND_GRAY)
    _add_para(doc, "Created product overview (problem, ICP, price, ascension path) and this build process document.")

    _add_heading(doc, "Platforms & Tools Used", level=1)
    _add_table(doc,
        ["Platform/Tool", "Purpose"],
        [
            ["Python 3.11", "Scoring model and Excel generation"],
            ["pandas", "Data manipulation and analysis"],
            ["openpyxl", "Excel file creation with formatting"],
            ["python-docx", "Professional Word document generation"],
            ["Claude (Anthropic)", "AI-assisted content generation, research"],
            ["VS Code + Claude Code", "Development environment"],
            ["ChatGPT (OpenAI)", "Custom GPT creation platform"],
            ["Git/GitHub", "Version control and code hosting"],
        ],
    )

    _add_heading(doc, "What Makes This Product Different", level=1)
    _add_numbered(doc, "Not a pitch deck: This is a real, functional product bundle — an Excel file that opens and works, a guide with actionable frameworks, and an interactive GPT.")
    _add_numbered(doc, "Built on real data: The 200-FO dataset was independently created and validated (Task 1), not scraped or generated synthetically.")
    _add_numbered(doc, "Clear ascension logic: Every component deliberately creates a natural next step toward PolarityIQ.")
    _add_numbered(doc, "Priced for velocity: At $197, the decision cost is trivial for the ICP.")

    out = ROOT / "products" / "docs" / "build_process.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Document 4: Custom GPT Prompt
# ---------------------------------------------------------------------------

def build_custom_gpt_doc() -> Path:
    doc = _setup_doc(
        title="Custom GPT\nSetup Guide",
        subtitle='FO Co-Investment Advisor — ChatGPT Custom GPT Configuration',
    )

    _add_heading(doc, "Setup Instructions", level=1)
    _add_numbered(doc, "Go to https://chat.openai.com → Explore GPTs → Create")
    _add_numbered(doc, 'Paste the System Instructions below into the "Instructions" field')
    _add_numbered(doc, "Upload co_investment_accelerator.xlsx as a knowledge file (critical — without it, the GPT has no proprietary data)")
    _add_numbered(doc, 'Name: "Co-Investment Intelligence Advisor"')
    _add_numbered(doc, 'Description: "Expert AI advisor for identifying and approaching family offices for co-investment partnerships. Powered by a scored database of 200 international family offices."')
    _add_numbered(doc, "Add the Conversation Starters listed below")

    _add_heading(doc, "System Instructions", level=1)
    # Read from the MD file to get the full system prompt
    md_path = ROOT / "products" / "custom_gpt_prompt.md"
    md_text = md_path.read_text(encoding="utf-8")
    # Extract the system instructions between ``` blocks
    start = md_text.find("```\nYou are the FO Co-Investment Advisor")
    end = md_text.find("```\n\n## Conversation Starters")
    if start >= 0 and end >= 0:
        system_text = md_text[start + 4:end].strip()
    else:
        system_text = "See custom_gpt_prompt.md for full system instructions."

    _add_code_block(doc, system_text)

    _add_heading(doc, "Conversation Starters", level=1)
    starters = [
        "Which family offices have the highest co-investment scores in Asia-Pacific?",
        "I'm raising a $50M growth equity fund focused on healthcare. Who should I approach first?",
        "Compare the co-investment readiness of European vs Middle Eastern family offices",
        "Show me all Green-readiness family offices that focus on technology",
    ]
    for s in starters:
        _add_numbered(doc, s)

    _add_heading(doc, "Example Conversations", level=1)

    _add_heading(doc, "Example 1: Healthcare Fund", level=2)
    _add_para(doc, "User: I'm raising a $100M Series B healthcare fund. Which family offices should I target?", bold=True)
    _add_para(doc, "The advisor searches the database and recommends:")
    _add_numbered(doc, "Poonawalla Family Office (India, SFO) — AUM: $8B, Co-Invest Score: 8/10, Accessibility: Green")
    _add_numbered(doc, "Builders Vision (US, SFO) — AUM: $5B, Co-Invest Score: 7/10, Accessibility: Green")
    _add_numbered(doc, "Dallah Albaraka Group (Saudi Arabia, SFO) — AUM: $8B, Co-Invest Score: 6/10, Accessibility: Yellow")

    _add_heading(doc, "Example 2: Small Raise Warning", level=2)
    _add_para(doc, "User: I'm raising a $5M seed round for my AI startup.", bold=True)
    _add_para(doc, "The advisor transparently notes that $5M is below most FO minimums and recommends angel networks, micro-VCs, or approaching FOs after securing institutional seed capital.")

    _add_heading(doc, "Example 3: Regional Comparison", level=2)
    _add_para(doc, "User: Compare European vs Middle Eastern family offices.", bold=True)
    _add_para(doc, "The advisor provides a detailed comparison covering average scores, readiness distribution, sector preferences, typical check sizes, and key hubs.")

    out = ROOT / "products" / "docs" / "custom_gpt_guide.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Document 5: Research Methodology
# ---------------------------------------------------------------------------

def build_methodology_doc() -> Path:
    doc = _setup_doc(
        title="Dataset Creation\nMethodology",
        subtitle="Task 1 — 200 International Family Office Intelligence Records",
    )

    _add_heading(doc, "Overview", level=1)
    _add_para(doc, "Created a validated dataset of 200 international family office intelligence records with 45 fields per record, covering 6 regions and 38 countries. The dataset prioritizes granularity and real-world accuracy over volume.")
    _add_para(doc, "Dataset creation date: March 2026", bold=True)

    _add_heading(doc, "Tech Stack", level=1)
    _add_table(doc,
        ["Tool", "Purpose"],
        [
            ["Python 3.11", "Primary programming language"],
            ["openpyxl", "XLSX reading/writing"],
            ["pandas", "Data manipulation, validation, statistics"],
            ["Claude (AI)", "Research assistance, data enrichment, pattern identification"],
            ["Web research", "Cross-referencing and validation against public sources"],
        ],
    )

    _add_heading(doc, "Data Acquisition Methodology", level=1)

    _add_heading(doc, "Seed List Construction", level=2)
    _add_para(doc, "Tier 1 — Well-documented FOs (Records 1-50):", bold=True)
    _add_para(doc, "Started with the most visible family offices — those with public filings (SEC Form ADV, 13F), press coverage, and known portfolios. Examples: Walton Enterprises, Cascade Investment, Bezos Expeditions, Soros Fund Management.")
    _add_para(doc, "Tier 2 — Regionally significant FOs (Records 51-150):", bold=True)
    _add_para(doc, "Expanded to internationally important family offices across Europe (Grosvenor, KIRKBI, Wendel, Exor), Asia-Pacific (Li Ka-Shing, Son Family, Premji Invest), Middle East (Olayan, Mansour, Dubai Holding), Latin America (Slim, Lemann, Santo Domingo), and Africa (Motsepe, Dangote, Sawiris).")
    _add_para(doc, "Tier 3 — Specialized/niche FOs (Records 151-200):", bold=True)
    _add_para(doc, "Added multi-family offices, impact-focused FOs, next-gen FOs, and historically significant families.")

    _add_heading(doc, "Source Hierarchy", level=2)
    _add_table(doc,
        ["Source", "Records Informed", "Data Types"],
        [
            ["SEC EDGAR (Form ADV, 13F)", "~50 US-based FOs", "Entity verification, AUM, holdings"],
            ["Forbes Global Billionaires", "~100 records", "Net worth, source of wealth"],
            ["LinkedIn", "~150 records", "Decision maker names/titles"],
            ["Bloomberg / Reuters", "~80 records", "Recent news, deal activity"],
            ["Crunchbase", "~60 records", "Portfolio companies, rounds"],
            ["Regulatory databases (FCA, MAS, FINMA)", "~30 records", "Regulatory status"],
            ["Conference records", "~40 records", "Conference attendance"],
            ["Corporate websites", "~120 records", "Website verification, emails"],
        ],
    )

    _add_heading(doc, "45-Field Schema", level=1)
    schema_sections = [
        ("Identity (7 fields)", "Family Office Name, Type (SFO/MFO), Founding Family, Year Established, HQ City, HQ Country, Region"),
        ("Contact (7 fields)", "Website, Primary DM LinkedIn, Email Pattern, Main Office Phone, LinkedIn Company URL, Contact Method, Conference Attendance"),
        ("People (4 fields)", "Primary Decision Maker, Primary DM Title, Secondary Decision Maker, Secondary DM Title"),
        ("Financials (2 fields)", "AUM ($B), Estimated Family Net Worth ($B)"),
        ("Operations (1 field)", "Employee Count (Est.)"),
        ("Investment (11 fields)", "Investment Strategy, Sector Focus, Geographic Focus, Asset Classes, Check Size Min/Max, Investment Stage, Direct Investment, Co-Invest Frequency, LP Status, GP/Direct Status"),
        ("Portfolio (1 field)", "Notable Portfolio Companies"),
        ("Relationships (3 fields)", "Fund Relationships, Co-Investor Relationships, ESG/Impact Level"),
        ("Signals (3 fields)", "Recent News (2024-2025), Recent Deals, Last Deal Date"),
        ("Governance (2 fields)", "Next-Gen Involvement, Philanthropy Focus"),
        ("Background (2 fields)", "Source of Wealth, Wealth Origin Sector"),
        ("Compliance (1 field)", "Regulatory Reference"),
        ("Meta (1 field)", "Data Confidence (High/Medium/Low)"),
    ]
    for section, fields in schema_sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{section}: ")
        run.font.bold = True
        run.font.color.rgb = BRAND_BLUE
        run = p.add_run(fields)

    _add_heading(doc, "Validation Methodology", level=1)

    _add_heading(doc, "AI Hallucination Safeguards", level=2)
    _add_para(doc, "Claude AI was used for research assistance and data enrichment, but all AI-generated data was validated against primary sources before inclusion:")
    _add_bullet(doc, "Every family office name was verified to exist as a real entity")
    _add_bullet(doc, "Decision maker names were cross-referenced against LinkedIn and press releases")
    _add_bullet(doc, "AUM figures were checked against Forbes wealth rankings and regulatory filings")
    _add_bullet(doc, "No AI-generated contact information was used without domain verification")
    _add_bullet(doc, "The Data Confidence field explicitly flags records where AI-assisted enrichment could not be fully validated")

    _add_heading(doc, "Data Quality Issues Found and Fixed", level=2)
    _add_bullet(doc, "Row 69 (Nair Pte Ltd): Complete column misalignment — all 45 fields reconstructed and realigned")
    _add_para(doc, "6 questionable entries replaced:", bold=True)
    replacements = [
        '"Mitsubishi Estate Family-Connected Office" → HAL Investments (Van der Vorm Family, Netherlands)',
        '"Mubadala-Adjacent Family Office" → Dallah Albaraka Group (Kamel Family, Saudi Arabia)',
        '"Wellcome Trust Family-Connected Office" → Reuben Brothers (Reuben Family, UK)',
        '"Alibaba Co-Founders Family Offices (Various)" → Yunfeng Capital (Jack Ma, Hong Kong)',
        '"Rabobank-linked Family Offices" → JAB Holding (Reimann Family, Luxembourg)',
        '"Wallenberg-Connected EQT Family Office" → Hoffmann Family Office (Roche, Switzerland)',
    ]
    for r in replacements:
        _add_bullet(doc, r)

    _add_heading(doc, "Confidence Scoring", level=2)
    _add_table(doc,
        ["Level", "Count", "Percentage", "Criteria"],
        [
            ["High", "116", "58%", "Multiple cross-referenced sources, verified decision makers"],
            ["Medium", "65", "32.5%", "Single source or partial verification"],
            ["Low", "19", "9.5%", "Limited public information, estimated fields"],
        ],
    )

    _add_heading(doc, "Key Insights", level=1)
    _add_numbered(doc, "SFO dominance: 185/200 (92.5%) are single-family offices.")
    _add_numbered(doc, "US concentration: 52/200 (26%) are US-based.")
    _add_numbered(doc, "AUM distribution: Highly skewed — median $15B vs. mean $28.2B.")
    _add_numbered(doc, "Co-investment trend: 69 FOs rated 'High' co-invest frequency, 73 'Medium'.")
    _add_numbered(doc, "ESG gap: 152/200 (76%) classified as 'Traditional/Not Disclosed'.")
    _add_numbered(doc, "Next-gen transition: 87 FOs multi-generational (3rd+), 75 founder-led, 35 have 2nd generation emerging.")

    _add_heading(doc, "Dataset Summary", level=1)
    _add_table(doc,
        ["Metric", "Value"],
        [
            ["Total Records", "200"],
            ["Total Fields", "45"],
            ["SFO/MFO Split", "185 / 15"],
            ["Regions Covered", "6"],
            ["Countries Covered", "38"],
            ["AUM Range", "$0.2B - $224.5B"],
            ["AUM Average", "$28.2B"],
            ["High Confidence", "116 (58%)"],
            ["Medium Confidence", "65 (32.5%)"],
            ["Low Confidence", "19 (9.5%)"],
        ],
    )

    out = ROOT / "research" / "docs" / "methodology.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Document 6: Conversion Analysis
# ---------------------------------------------------------------------------

def build_conversion_analysis() -> Path:
    doc = _setup_doc(
        title="PolarityIQ\nConversion Analysis",
        subtitle="Free Trial → Paid User Conversion Optimization Plan",
    )

    _add_heading(doc, "Executive Summary", level=1)
    _add_para(doc, "After a thorough review of all three Falcon Scaling / PolarityIQ web properties (app.polarityiq.com, polarityiq.com, falconscaling.com), this analysis identified 3 critical trust-breaking inconsistencies, 5 high-impact conversion friction points, and 12 specific recommendations organized by timeline.")
    _add_callout(doc, "The platform's core value proposition is strong — 10,000+ FO profiles with verified contact data at $295-$1,995/mo vs. $15K-$25K/yr from legacy providers.", label="Key Finding")

    _add_heading(doc, "1. Platform Audit — What Was Observed", level=1)

    _add_heading(doc, "1.1 app.polarityiq.com (SaaS Platform)", level=2)
    _add_para(doc, "What works well:", bold=True)
    _add_bullet(doc, 'Strong hero section: "Comprehensive Family Office Intelligence" with social proof')
    _add_bullet(doc, "Comparison table (PolarityIQ vs Legacy Providers) is the single strongest conversion element")
    _add_bullet(doc, "10 specific user personas help visitors self-identify immediately")
    _add_bullet(doc, "3-step How It Works with visual mockups is clean and effective")
    _add_bullet(doc, "Platform Guide page acts as an interactive demo substitute")

    _add_para(doc, "Pricing structure observed:", bold=True)
    _add_table(doc,
        ["Plan", "Price", "Tokens", "Per Token"],
        [
            ["Starter", "$295/mo", "100", "$2.95"],
            ["Basic (Most Popular)", "$395/mo", "250", "$1.58"],
            ["Pro", "$795/mo", "600", "$1.33"],
            ["Premium", "$1,995/mo", "2,000", "$1.00"],
        ],
    )
    _add_para(doc, "Trial: 7 days, 10 tokens (contact reveals), credit card required.", italic=True, color=BRAND_GRAY)

    _add_heading(doc, "2. Critical Trust-Breaking Issues (Fix Immediately)", level=1)

    issues = [
        ("Issue #1: Data Count Contradiction",
         'The H1 subtext says "10,000+ Family Offices" but the stats bar says "800+ Family Offices."',
         'Reconcile to one number: "10,000+ Family Office profiles | 800+ with verified decision-maker contacts."'),
        ("Issue #2: Credit Card Requirement Contradiction",
         'The guide page CTA says "No credit card required" while the signup page states "Credit card required."',
         "Update the guide page CTA to match reality, or remove the CC requirement."),
        ("Issue #3: 'Not a SaaS Platform' Claim",
         'polarityiq.com states "Polarity IQ is not a SaaS platform" while app.polarityiq.com IS a SaaS platform.',
         "Rewrite to distinguish the custom data business from the SaaS platform."),
    ]
    for title, problem, fix in issues:
        _add_heading(doc, title, level=2)
        _add_para(doc, f"Problem: {problem}")
        _add_para(doc, f"Fix: {fix}", bold=True)

    _add_heading(doc, "3. High-Impact Conversion Friction Points", level=1)

    frictions = [
        ("Friction #1: Signup Flow — No SSO, All Fields on One Page",
         "Add LinkedIn/Google SSO. Progressive wizard: Step 1 = Email → Step 2 = Name + Password → Step 3 = Plan selection."),
        ("Friction #2: Deferred Plan Selection Creates Uncertainty",
         "Show all 4 plan cards during signup with trial framing. Highlight per-token cost difference."),
        ("Friction #3: Token Allocation May Not Reach 'Aha Moment'",
         "A/B test token allocation (10 vs 15 vs 25). Consider bonus tokens for actions (+5 for first search, +5 for saving, +5 for export)."),
        ("Friction #4: No Onboarding Email Sequence",
         "Implement 5-email activation sequence: Day 0 (marquee FO data), Day 1 (token tutorial), Day 3 (personalized), Day 5 (urgency), Day 6 (final push)."),
        ("Friction #5: Two Disconnected Web Properties",
         "Add persistent cross-site navigation bar. Unify brand narrative. Cross-promote between SaaS and custom data."),
    ]
    for title, rec in frictions:
        _add_heading(doc, title, level=2)
        _add_para(doc, f"Recommendation: {rec}")

    _add_heading(doc, "4. What's Working — Don't Change These", level=1)
    _add_numbered(doc, "Comparison table (PolarityIQ vs Legacy Providers) — most powerful conversion element")
    _add_numbered(doc, "10 user personas — specific ICP targeting helps visitors self-identify")
    _add_numbered(doc, "Triple CTA (Free Trial / Guide / Demo) — caters to different buyer intent")
    _add_numbered(doc, 'Token reveal system — the "blur → reveal" UX is elegant')
    _add_numbered(doc, "Platform Guide acts as an interactive demo without requiring signup")
    _add_numbered(doc, '"A+/A/B" email quality badges — unique differentiator')
    _add_numbered(doc, "Logo bar (Saison Capital, Cresset, Hamilton Lane) — institutional credibility")
    _add_numbered(doc, '"$0 due today" messaging on signup — strong reassurance')

    _add_heading(doc, "5. Competitive Context", level=1)
    _add_table(doc,
        ["Feature", "PolarityIQ", "Fintrx", "Dakota", "PitchBook"],
        [
            ["FO-specific database", "10,000+ profiles", "3,800+ firms", "Partial", "Partial"],
            ["Verified emails", "Yes (A+/A/B)", "Yes", "Limited", "Yes"],
            ["Price", "$295-$1,995/mo", "~$10K-15K/yr", "~$5K-10K/yr", "~$20K-40K/yr"],
            ["Free trial", "7 days, 10 tokens", "No", "No", "No"],
            ["Self-serve signup", "Yes", "No (sales-led)", "No", "No"],
        ],
    )
    _add_callout(doc, "PolarityIQ is the ONLY platform where a fund manager can sign up, search, and reveal a FO decision-maker's email in 5 minutes, without talking to anyone.", label="Key Advantage")

    _add_heading(doc, "6. Recommendations by Timeline", level=1)

    _add_heading(doc, "Quick Wins (1-2 Weeks)", level=2)
    _add_numbered(doc, "Fix the three data contradictions — pure copy fixes")
    _add_numbered(doc, "Add competitive pricing anchor to pricing section")
    _add_numbered(doc, "A/B test trial token allocation (10 vs 15 vs 20)")
    _add_numbered(doc, "Optimize token exhaustion modal with specific opportunity cost")
    _add_numbered(doc, "Add password strength indicator to signup")

    _add_heading(doc, "Medium-Term (2-6 Weeks)", level=2)
    _add_numbered(doc, "Redesign signup as multi-step wizard")
    _add_numbered(doc, "Build 5-email activation sequence with real FO data")
    _add_numbered(doc, "Add LinkedIn SSO (3-4 weeks engineering)")
    _add_numbered(doc, "Implement trial win-back sequence (Day 8, 14, 30)")
    _add_numbered(doc, "Add in-app live chat for high-intent trial users")

    _add_heading(doc, "Strategic (1-3 Months)", level=2)
    _add_numbered(doc, "Unify navigation between app.polarityiq.com and polarityiq.com")
    _add_numbered(doc, "Add annual pricing option (validate first)")
    _add_numbered(doc, "Retire or redirect falconscaling.com")
    _add_numbered(doc, "Build graduated paywall — let free users search, gate contact reveals")
    _add_numbered(doc, "Sales-assist for Premium tier ($1,995/mo = $24K/yr ACV)")

    _add_heading(doc, "7. Measurement Framework", level=1)
    _add_table(doc,
        ["Metric", "Definition", "B2B SaaS Benchmark", "Target"],
        [
            ["Signup completion", "% of pricing page → signup", "20-40%", "Improve 20%+"],
            ["Trial activation", "% signups who reveal >= 1 contact", "40-60%", ">70%"],
            ["Trial → Paid", "% of trial users who subscribe", "5-8%", ">10%"],
            ["Time to first reveal", "Seconds from signup to first use", "N/A", "<180s"],
            ["Monthly churn", "% paid users cancelling/month", "5-7%", "<6%"],
            ["Upgrade rate", "% Starter → Basic+", "10-15%", ">15%"],
        ],
    )

    _add_callout(doc, "Leading indicator of conversion: 'User reveals 5+ contacts AND exports at least 1 CSV within the first 4 days.' Track this cohort and optimize onboarding toward it.", label="Activation Metric")

    out = ROOT / "analysis" / "docs" / "conversion_analysis.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("Professional Document Generator")
    print("=" * 60)

    docs = [
        ("Approach Guide", build_approach_guide),
        ("Product Overview", build_product_overview),
        ("Build Process", build_build_process),
        ("Custom GPT Guide", build_custom_gpt_doc),
        ("Research Methodology", build_methodology_doc),
        ("Conversion Analysis", build_conversion_analysis),
    ]

    for name, builder in docs:
        print(f"\n  Generating {name}...")
        path = builder()
        print(f"    Saved: {path}")

    print(f"\n{'=' * 60}")
    print("All 6 professional DOCX documents generated successfully.")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
